import os
from docx import Document
from docx.opc.exceptions import PackageNotFoundError
from pypdf import PdfReader
from fastapi import UploadFile



def read_document_content(file:UploadFile,file_extension:str) -> str:
    """
    Reads and returns the text content from a .txt, .pdf, or .docx file.
    """
    if not file:
        return f"Error: file is empty"

    if file_extension == '.txt':
        return _parse_txt(file)
    elif file_extension == '.pdf':
        return _parse_pdf(file)
    elif file_extension == '.docx':
        return _parse_docx(file)
    else:
        return f"Error: Unsupported file type: {file_extension}"


def _parse_txt(file:UploadFile) -> str:
    try:
        contents = file.read()
        return contents.decode('utf-8')
    except UnicodeDecodeError:
        return f"Error: Unable to decode TXT file {file.filename}. Please ensure it is a valid UTF-8 encoded text file."
    except Exception as e:
        return f"Error: reading {file.filename} failed with error: {e}"

def _parse_pdf(file: UploadFile) -> str:
    text_content = ""
    try:
        reader = PdfReader(file.file)
        for page in reader.pages:
            text_content += page.extract_text() + "\n"
        return text_content
    except Exception as e:
        return f"Error: reading PDF file {file.filename}: {e}"

def _parse_docx(file:UploadFile) -> str:
    text_content = []
    try:
        document = Document(file.file)
        for paragraph in document.paragraphs:
            text_content.append(paragraph.text)
        for table in document.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    row_text.append(cell.text)
                text_content.append(" ".join(row_text)) # Using space for simplicity in tables
        return "\n".join(text_content)
    except PackageNotFoundError:
        return f"Error: DOCX file invalid at {file.filename}"
    except Exception as e:
        return f"Error: reading DOCX file {file.filename}: {e}"